# -*- coding: utf-8 -*-
"""
============================================================
MODULE: exporters/word_khbd.py
Nhiệm vụ: Động cơ kết xuất Word chuẩn 5512, tích hợp ScienceNormalizer, 
xử lý LaTeX và chống vỡ khung bảng biểu.
============================================================
"""

import io
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from export.markdown_tokenizer import MarkdownTokenizer
except ImportError:
    MarkdownTokenizer = None

class ScienceNormalizer:
    """Bộ lọc chuẩn hóa khoa học (Toán, Lý, Hóa) chống cắt cụt dấu căn và ký hiệu"""
    SUB = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
    SUP = str.maketrans("0123456789+-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻")
    MAP = {
        r'\perp': '⊥', r'\circ': '°', r'\ne': '≠', r'\le': '≤', r'\ge': '≥', 
        r'\times': '×', r'\div': '÷', r'\triangle': '△', r'\angle': '∠', 
        r'\rightarrow': '→', r'\Rightarrow': '⇒', r'\approx': '≈',
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\pi': 'π', 
        r'\sum': '∑', r'\int': '∫'
    }

    @classmethod
    def normalize(cls, text: str) -> str:
        if not text: 
            return ""
        text = str(text).replace('$', '').strip()
        text = re.sub(r'\\sqrt\s*\{([\s\S]+?)\}', r'√( \1 )', text)
        text = re.sub(r'\\sqrt\s*([a-zA-Z0-9]+)', r'√\1', text)
        while r'\frac{' in text:
            text = re.sub(r'\\frac\{([\s\S]+?)\}\{([\s\S]+?)\}', r'( \1 / \2 )', text)
        text = re.sub(r'([A-Z][a-z]?|\))(\d+)', lambda m: m.group(1) + m.group(2).translate(cls.SUB), text)
        text = re.sub(r'([A-Za-z₀₁₂₃₄₅₆₇₈₉\)]+)\^(\d*[+\-])', lambda m: m.group(1) + m.group(2).translate(cls.SUP), text)
        for k, v in cls.MAP.items(): 
            text = text.replace(k, v)
        return re.sub(r'\\text\{([\s\S]+?)\}', r'\1', text)

class WordExportEngine:

    @staticmethod
    def _set_font(run, font_name="Times New Roman"):
        try:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass

    @staticmethod
    def _set_shading(p_or_cell, color_hex):
        element = p_or_cell._element.get_or_add_pPr() if hasattr(p_or_cell, 'paragraphs') else p_or_cell._element.get_or_add_tcPr()
        shd = element.find(qn("w:shd"))
        if shd is None: 
            shd = OxmlElement("w:shd")
            element.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color_hex)

    @classmethod
    def _render_inline(cls, p, tokens: List[Dict[str, Any]]):
        if not tokens: 
            return
        for t in tokens:
            tt = t.get("type")
            c = t.get("content", "") or t.get("text", "")
            
            if tt in ["text", "bold", "italic", "underline", "strike", "highlight", "inline_code"]:
                run = p.add_run(c)
                cls._set_font(run, "Courier New" if tt == "inline_code" else "Times New Roman")
                run.font.size = Pt(13)
                if tt == "bold": run.bold = True
                elif tt == "italic": run.italic = True
                elif tt == "underline": run.underline = True
                elif tt == "inline_code": 
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(199, 37, 78)
                    
            elif tt in ["inline_math", "math_inline", "math", "block_math"]:
                normalized_math = ScienceNormalizer.normalize(c)
                run = p.add_run(normalized_math)
                run.font.italic = True
                cls._set_font(run, "Times New Roman")
            else:
                run = p.add_run(str(c))
                cls._set_font(run, "Times New Roman")
                run.font.size = Pt(13)

    @classmethod
    def convert_markdown_to_docx_bytes(cls, markdown_text: str) -> bytes:
        doc = Document()
        
        # Thiết lập chuẩn trang A4
        for s in doc.sections:
            s.page_height, s.page_width = Inches(11.69), Inches(8.27)
            s.top_margin, s.bottom_margin, s.right_margin = Inches(0.79), Inches(0.79), Inches(0.79)
            s.left_margin = Inches(1.18)

        # Chuyển đổi định dạng LaTeX \[ \] và \( \) sang chuẩn $$ và $ trước khi Tokenizer đọc
        markdown_text = re.sub(r'\\\[(.*?)\\\]', r'$$\1$$', markdown_text, flags=re.DOTALL)
        markdown_text = re.sub(r'\\\((.*?)\\\)', r'$\1$', markdown_text)

        if MarkdownTokenizer and hasattr(MarkdownTokenizer, 'parse'):
            try:
                ast_nodes = MarkdownTokenizer.parse(markdown_text)
            except Exception:
                ast_nodes = [{"type": "paragraph", "tokens": [{"type": "text", "content": markdown_text}]}]
        else:
            ast_nodes = [{"type": "paragraph", "tokens": [{"type": "text", "content": markdown_text}]}]

        for node in ast_nodes:
            nt = node.get("type")
            
            if nt == "paragraph":
                p = doc.add_paragraph()
                cls._render_inline(p, node.get("tokens", []))
                
            elif nt == "heading":
                lv = min(max(node.get("level", 1), 1), 3)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                if lv == 1: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cls._render_inline(p, node.get("tokens", []))
                for r in p.runs: 
                    r.bold = True
                    r.font.size = Pt(17 - lv)
                    cls._set_font(r)
                    
            elif nt == "list_item":
                st = 'List Number' if node.get("style") == "number" else 'List Bullet'
                p = doc.add_paragraph(style=st)
                p.paragraph_format.left_indent = Inches(0.25 * node.get("level", 1) + 0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                cls._render_inline(p, node.get("tokens", []))
                
            elif nt == "table":
                rows, headers, cols = node.get("rows", []), node.get("headers", []), node.get("cols", 1)
                if cols > 0:
                    table = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=cols)
                    table.style = 'Table Grid'
                    table.alignment = 1
                    
                    # BẮT BUỘC: Khóa Autofit để bảng không bị vỡ như kỹ sư chỉ định
                    table.autofit = False
                    col_width = Inches(6.3 / cols)

                    r_idx = 0
                    if headers:
                        h_row = table.rows[0]
                        h_row._element.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
                        for c_idx, cell_n in enumerate(headers):
                            cell = h_row.cells[c_idx]
                            cell.width = col_width # Khóa cứng độ rộng
                            cls._set_shading(cell, "EAEAEA")
                            p = cell.paragraphs[0]
                            cls._render_inline(p, cell_n.get("content", []))
                            for r in p.runs: r.bold = True
                        r_idx = 1
                        
                    for loop_idx, r_data in enumerate(rows):
                        row = table.rows[r_idx]
                        row._element.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
                        bg = "F9F9F9" if loop_idx % 2 == 1 else "FFFFFF"
                        for c_idx, cell_n in enumerate(r_data):
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.width = col_width # Khóa cứng độ rộng
                                cls._set_shading(cell, bg)
                                cls._render_inline(cell.paragraphs[0], cell_n.get("content", []))
                        r_idx += 1

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

def export_word(data_cache) -> bytes:
    try:
        md_text = data_cache.get("ai_generated_content", "") if isinstance(data_cache, dict) else str(data_cache)
        return WordExportEngine.convert_markdown_to_docx_bytes(md_text)
    except Exception as e:
        fallback_doc = Document()
        fallback_doc.add_paragraph(f"LỖI KẾT XUẤT: {str(e)}")
        bio = io.BytesIO()
        fallback_doc.save(bio)
        return bio.getvalue()

class KhbdWordExporter:
    @classmethod
    def export_to_word(cls, data_cache) -> bytes:
        return export_word(data_cache)
